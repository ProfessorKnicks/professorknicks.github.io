import random
from docx import Document
from docx.shared import Pt
from docx.enum.table import WD_TABLE_ALIGNMENT

options = [
"Blinding Lights",
"Heat Waves",
"As It Was",
"STAY",
"Shape of You",
"Levitating",
"Despacito",
"Mr. Brightside",
"Lose Yourself",
"Bad guy",
"We Belong Together",
"Uptown Funk",
"Happy",
"Old Town Road",
"Good 4 U",
"Viva La Vida",
"Thrift Shop",
"Party Rock Anthem",
"Big Girls Dont Cry",
"Somebody That I Used To Know",
"All of Me",
"Shallow",
"Say My Name",
"Crazy",
"Counting Stars",
"Hey There Delilah",
"Low",
"Dont Stop The Music",
"Chandelier",
"See You Again",
"I Kissed A Girl",
"Gods Plan",
"Wake Me Up",
"Cheerleader",
"Rolling in the Deep",
"Hey, Soul Sister",
"Firework",
"Umbrella",
"Titanium",
"Hips Dont Lie",
"On The Floor",
"I Want It That Way",
"Dynamite",
"Stereo Hearts",
"Diamonds",
"What Makes You Beautiful"
]

cards = 2

# Create a new Word document
doc = Document()

# Set font size for bingo card title
font_title = doc.styles['Title']
font_title.font.size = Pt(20)

# Loop through each card
for i in range(cards):
    # Add a new page for each card
    if i != 0:
        doc.add_page_break()

    # Add the card number
    doc.add_paragraph("Pop Musical Bingo", style='Title').alignment = WD_TABLE_ALIGNMENT.CENTER

    # Create a new table with a 5x5 grid
    table = doc.add_table(rows=5, cols=5)

    # Set table style
    table.style = 'Table Grid'

    # Get the width of the page
    page_width = doc.sections[0].page_width - doc.sections[0].left_margin - doc.sections[0].right_margin

    # Adjust table width to fill the page
    table.width = page_width

    # Adjust cell width and height to make them square
    cell_width = page_width / 5
    cell_height = cell_width
    for row in table.rows:
        for cell in row.cells:
            cell.width = cell_width
            cell.height = cell_height
            # Center the text vertically and horizontally
            cell.vertical_alignment = WD_TABLE_ALIGNMENT.CENTER
            cell.paragraphs[0].alignment = WD_TABLE_ALIGNMENT.CENTER
            # Set font size for options
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(16)

    # Add the options to the table
    for j in range(5):
        for k in range(5):
            num = random.choice(options)
            while num in [cell.text for row in table.rows for cell in row.cells]:
                num = random.choice(options)
            cell = table.cell(j, k)
            cell.text = num

# Save the Word document
doc.save("bingo_cards.docx")
